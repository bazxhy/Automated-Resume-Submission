"""
简历解析模块 — 支持 PDF / DOCX / TXT 格式
提取：姓名、联系方式、技能标签、工作经历、教育背景
"""

from __future__ import annotations

import re
from dataclasses import dataclass, field
from pathlib import Path
from typing import Optional

from utils import resolve_path, clean_text, get_logger

logger = get_logger("resume")


# ---------- 数据结构 ----------

@dataclass
class Experience:
    """单段工作/教育经历"""
    company_or_school: str = ""
    title_or_major: str = ""
    start_date: str = ""
    end_date: str = ""
    description: str = ""


@dataclass
class ResumeData:
    """解析后的简历数据结构"""
    name: str = ""
    phone: str = ""
    email: str = ""
    skills: list[str] = field(default_factory=list)
    work_experiences: list[Experience] = field(default_factory=list)
    education: list[Experience] = field(default_factory=list)
    raw_text: str = ""
    parsed_from: str = ""

    def to_dict(self) -> dict:
        return {
            "name": self.name,
            "phone": self.phone,
            "email": self.email,
            "skills": self.skills,
            "work_experiences": [
                {"company": e.company_or_school, "title": e.title_or_major,
                 "period": f"{e.start_date} - {e.end_date}",
                 "description": e.description}
                for e in self.work_experiences
            ],
            "education": [
                {"school": e.company_or_school, "major": e.title_or_major,
                 "period": f"{e.start_date} - {e.end_date}"}
                for e in self.education
            ],
        }


# ---------- 解析器 ----------

class ResumeParser:
    """简历解析器"""

    # 技能词库（英文 + 中文 + 框架 + 嵌入式）
    SKILL_KEYWORDS = [
        # 编程语言
        "Python", "Java", "Go", "Golang", "C++", "C#", "PHP", "Ruby",
        "JavaScript", "TypeScript", "Kotlin", "Swift", "Rust", "Scala",
        "Shell", "Bash", "C", "汇编", "Assembly", "Verilog", "VHDL", "MATLAB",
        # 前端
        "React", "Vue", "Angular", "HTML5", "CSS3", "jQuery", "Bootstrap",
        "Webpack", "Vite", "微信小程序", "小程序",
        # 后端框架
        "Django", "Flask", "FastAPI", "Spring", "Spring Boot", "Spring Cloud",
        "MyBatis", "Hibernate", "Express", "NestJS", "Gin", "Beego",
        # 数据库
        "MySQL", "PostgreSQL", "MongoDB", "Redis", "Oracle", "SQL Server",
        "Elasticsearch", "Memcached", "HBase", "Hive", "ClickHouse",
        # 中间件 & 消息队列
        "Kafka", "RabbitMQ", "RocketMQ", "Zookeeper", "Nginx", "Tomcat",
        "Dubbo", "Nacos", "Sentinel", "Gateway",
        # 运维 & 容器
        "Docker", "Kubernetes", "K8s", "Jenkins", "GitLab CI", "GitHub Actions",
        "Terraform", "Ansible", "Prometheus", "Grafana",
        # 基础
        "Linux", "Git", "HTTP", "TCP/IP", "WebSocket", "RESTful", "GraphQL", "gRPC",
        # 大数据 & AI
        "Hadoop", "Spark", "Flink", "Storm", "TensorFlow", "PyTorch",
        "Scikit-learn", "Pandas", "NumPy", "Matplotlib", "Jupyter",
        # 云服务
        "AWS", "Azure", "阿里云", "腾讯云", "华为云",
        # 软技能 & 领域
        "微服务", "分布式", "高并发", "多线程", "爬虫", "数据分析", "机器学习",
        "敏捷开发", "CI/CD",
        # 中文常见描述
        "数据库", "后端开发", "前端开发", "全栈", "运维", "测试", "自动化",
        # === 嵌入式 ===
        "嵌入式", "单片机", "ARM", "STM32", "GD32", "ESP32", "AVR", "PIC",
        "FPGA", "DSP", "Zynq", "Xilinx", "Altera",
        "RTOS", "FreeRTOS", "uCOS", "ThreadX", "RT-Thread", "VxWorks",
        "Linux驱动", "Linux内核", "BSP", "设备树", "Bootloader", "U-Boot",
        "I2C", "SPI", "UART", "USART", "CAN", "CAN-FD", "RS232", "RS485",
        "USB", "PCIe", "Ethernet", "SDIO", "PWM", "ADC", "DAC", "GPIO",
        "ZigBee", "BLE", "蓝牙", "WiFi", "LoRa", "NB-IoT", "4G", "5G",
        "Keil", "IAR", "GCC", "Makefile", "CMake", "交叉编译", "烧录",
        "示波器", "万用表", "逻辑分析仪", "硬件调试", "原理图", "PCB",
        "传感器", "电机控制", "电源管理", "FOC", "编码器",
        "物联网", "IoT", "智能家居", "工业控制", "汽车电子",
    ]

    # 章节关键词（支持多种写法）
    SECTION_KEYWORDS = {
        "work": [
            "工作经历", "工作经验", "工作经历", "项目经验",
            "项目经历", "实习经历", "社会实践", "Work Experience",
            "WORK EXPERIENCE", "PROJECTS",
        ],
        "edu": [
            "教育背景", "教育经历", "教育", "学历", "学习经历",
            "Education", "EDUCATION",
        ],
        "skills_section": [
            "技能", "专业技能", "技术栈", "技术能力", "个人技能",
            "Skills", "SKILLS", "Tech Stack",
        ],
    }

    def __init__(self, file_path: str):
        self.file_path = resolve_path(file_path)
        self.logger = logger

    def parse(self) -> ResumeData:
        """解析简历，返回结构化数据"""
        if not self.file_path.exists():
            raise FileNotFoundError(f"简历文件不存在: {self.file_path}")

        ext = self.file_path.suffix.lower()
        raw_text = self._read_file(ext)
        self.logger.info(f"原文长度: {len(raw_text)} 字符")

        resume = ResumeData(raw_text=raw_text, parsed_from=self.file_path.name)

        self._extract_contact(resume)
        self._extract_skills(resume)
        self._extract_sections(resume)

        # 如果结构化提取的工作经历为空，走全文本兜底
        if len(resume.work_experiences) == 0:
            self.logger.info("结构化提取工作经历为 0，尝试全局扫描...")
            self._parse_work_experiences(resume, raw_text.split("\n"))

        self.logger.info(
            f"解析完成: {resume.name or '未知姓名'} | "
            f"技能 {len(resume.skills)} 项 | "
            f"工作经历 {len(resume.work_experiences)} 段 | "
            f"教育 {len(resume.education)} 段"
        )
        return resume

    def _read_file(self, ext: str) -> str:
        """根据后缀读取文件内容，取提取质量最好的方式"""
        if ext == ".pdf":
            return self._read_pdf_best()
        elif ext == ".docx":
            return self._read_docx()
        elif ext == ".txt":
            return self._read_txt()
        else:
            raise ValueError(f"不支持的简历格式: {ext}")

    def _read_pdf_best(self) -> str:
        """PDF: 同时尝试 pdfplumber 和 PyMuPDF，返回文本更长的那个"""
        text_plumber = ""
        text_mupdf = ""

        # 方式 1: pdfplumber
        try:
            import pdfplumber
            with pdfplumber.open(str(self.file_path)) as pdf:
                pages = [page.extract_text() or "" for page in pdf.pages]
            text_plumber = "\n".join(pages)
            self.logger.debug(f"pdfplumber: {len(text_plumber)} 字符")
        except Exception as e:
            self.logger.debug(f"pdfplumber 失败: {e}")

        # 方式 2: PyMuPDF
        try:
            import fitz
            doc = fitz.open(str(self.file_path))
            text_mupdf = "\n".join(page.get_text() for page in doc)
            doc.close()
            self.logger.debug(f"PyMuPDF: {len(text_mupdf)} 字符")
        except Exception as e:
            self.logger.debug(f"PyMuPDF 失败: {e}")

        # 取文本量更大的
        if len(text_mupdf) > len(text_plumber):
            self.logger.info(f"使用 PyMuPDF 提取结果 ({len(text_mupdf)} 字符)")
            return text_mupdf
        elif len(text_plumber) > 0:
            self.logger.info(f"使用 pdfplumber 提取结果 ({len(text_plumber)} 字符)")
            return text_plumber
        else:
            self.logger.warning("PDF 文本提取为空！可能是图片型 PDF，建议转换为文本或使用附件中的 TXT/DOCX 版本")
            return ""

    def _read_docx(self) -> str:
        """读取 DOCX"""
        from docx import Document
        doc = Document(str(self.file_path))
        lines = [p.text for p in doc.paragraphs]
        # 也读表格中的内容
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    lines.append(cell.text)
        return "\n".join(lines)

    def _read_txt(self) -> str:
        """读取 TXT（自动检测编码）"""
        encodings = ["utf-8", "gbk", "gb2312", "gb18030", "utf-16"]
        for enc in encodings:
            try:
                with open(self.file_path, "r", encoding=enc) as f:
                    return f.read()
            except (UnicodeDecodeError, UnicodeError):
                continue
        with open(self.file_path, "rb") as f:
            return f.read().decode("utf-8", errors="ignore")

    # ---------- 信息提取 ----------

    def _extract_contact(self, resume: ResumeData):
        """提取姓名、手机号、邮箱"""
        text = resume.raw_text

        # 手机号
        phone_match = re.search(r"1[3-9]\d{9}", text)
        if phone_match:
            resume.phone = phone_match.group()

        # 邮箱
        email_match = re.search(r"[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}", text)
        if email_match:
            resume.email = email_match.group()

        # 姓名：扫描前 20 行非空文本
        lines = [l.strip() for l in text.split("\n") if l.strip()]
        for line in lines[:20]:
            # 显式 "姓名: xxx" / "Name: xxx"
            m = re.search(r"(?:姓名|Name)[：:]\s*(\S{2,4})", line)
            if m:
                resume.name = m.group(1)
                return
        # 兜底：第 1~3 行中纯中文且不是标题词
        skip_words = ["简历", "resume", "cv", "个人简介", "求职", "应聘", "联系", "电话", "邮箱"]
        for line in lines[:5]:
            clean = line.strip()
            # 2-4 个中文字符、不含符号
            if 2 <= len(clean) <= 4 and re.fullmatch(r"[\u4e00-\u9fff·]+", clean):
                if not any(sk in clean.lower() for sk in skip_words):
                    resume.name = clean
                    return

    def _extract_skills(self, resume: ResumeData):
        """从全文匹配技能关键词"""
        text_lower = resume.raw_text.lower()
        matched = set()
        for skill in self.SKILL_KEYWORDS:
            if skill.lower() in text_lower:
                matched.add(skill)
        resume.skills = sorted(matched, key=lambda s: self.SKILL_KEYWORDS.index(s) if s in self.SKILL_KEYWORDS else 999)

    def _extract_sections(self, resume: ResumeData):
        """按章节关键词分割文本"""
        text = resume.raw_text
        lines = text.split("\n")

        current_section = None
        work_lines = []
        edu_lines = []

        for line in lines:
            stripped = line.strip()
            if not stripped:
                continue

            # 检测章节切换
            for section, keywords in self.SECTION_KEYWORDS.items():
                if section in ("skills_section",):
                    continue  # 技能单独处理
                for kw in keywords:
                    if kw in stripped:
                        current_section = section
                        break
                if current_section == section:
                    break

            if current_section == "work":
                work_lines.append(stripped)
            elif current_section == "edu":
                edu_lines.append(stripped)

        self._parse_work_experiences(resume, work_lines)
        self._parse_education(resume, edu_lines)

    def _extract_sections_simple(self, resume: ResumeData):
        """简单文本的章节提取（TXT 回退方案）"""
        self._extract_sections(resume)

    def _parse_work_experiences(self, resume: ResumeData, lines: list[str]):
        """从行列表中解析工作经历"""
        if not lines:
            return

        # 更宽松的日期匹配
        # 匹配: 2020.01-2022.06 / 2020/01-至今 / 2020.01 - 2022.06 / 2020年01月-至今
        date_pattern = re.compile(
            r"(\d{4}\s*[年./-]?\s*\d{1,2}\s*[月]?)\s*"
            r"[-~–—至到]\s*"
            r"(\d{4}\s*[年./-]?\s*\d{1,2}\s*[月]?|至今|现在|Now|now)"
        )

        current = Experience()
        for line in lines:
            dm = date_pattern.search(line)
            if dm:
                # 保存上一段
                if current.company_or_school or current.title_or_major:
                    resume.work_experiences.append(current)
                current = Experience(
                    start_date=dm.group(1).strip(),
                    end_date=dm.group(2).strip(),
                )
                # 日期前面的文字 = 公司名 + 职位
                before_date = line[:dm.start()].strip()
                current.company_or_school = clean_text(before_date) if before_date else ""
            elif current.company_or_school or current.title_or_major:
                # 日期后面的补充行
                stripped = clean_text(line)
                if not current.title_or_major and len(stripped) > 0:
                    current.title_or_major = stripped
                else:
                    if stripped:
                        current.description += stripped + " "

        if current.company_or_school or current.title_or_major:
            resume.work_experiences.append(current)

        # 给没有 title 的填充
        for exp in resume.work_experiences:
            if not exp.title_or_major:
                # 尝试从描述中提取
                for keyword in ["工程师", "开发", "经理", "主管", "实习", "专员", "负责人", "技术总监"]:
                    if keyword in exp.description or keyword in exp.company_or_school:
                        exp.title_or_major = keyword
                        break
            if not exp.description:
                exp.description = ""

    def _parse_education(self, resume: ResumeData, lines: list[str]):
        """解析教育背景"""
        if not lines:
            return

        date_pattern = re.compile(
            r"(\d{4}\s*[年./-]?\s*\d{1,2}\s*[月]?)\s*"
            r"[-~–—至到]\s*"
            r"(\d{4}\s*[年./-]?\s*\d{1,2}\s*[月]?|至今|现在)"
        )
        degree_keywords = ["本科", "硕士", "博士", "大专", "研究生", "学士", "MBA", "高中"]

        current = Experience()
        for line in lines:
            stripped = clean_text(line)
            if not stripped:
                continue

            dm = date_pattern.search(line)
            if dm:
                if current.company_or_school:
                    resume.education.append(current)
                current = Experience(
                    start_date=dm.group(1).strip(),
                    end_date=dm.group(2).strip(),
                )
                current.company_or_school = clean_text(line)
                # 也尝试提取学位
                for dk in degree_keywords:
                    if dk in line:
                        current.title_or_major = dk
                        break
            elif current.company_or_school:
                if any(dk in stripped for dk in degree_keywords) and not current.title_or_major:
                    current.title_or_major = stripped
                else:
                    current.description += stripped + " "

        if current.company_or_school:
            resume.education.append(current)


# ---------- 便捷入口 ----------

def parse_resume(file_path: str) -> ResumeData:
    """一键解析简历"""
    parser = ResumeParser(file_path)
    return parser.parse()
